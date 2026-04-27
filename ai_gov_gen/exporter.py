"""DOCX and PDF export engine for AI Gov Gen.

This module converts generated governance artifacts into downloadable
DOCX files (via python-docx) and PDF files (via WeasyPrint).

The module exposes two primary export functions:

* :func:`export_docx` — writes a ``GovernanceChecklist``, list of
  ``SSPEntry`` objects, and a ``PolicyDocument`` to a single formatted
  ``.docx`` file using python-docx.

* :func:`export_pdf` — renders the policy HTML produced by
  :func:`~ai_gov_gen.generator.render_policy_html` to a PDF byte stream
  using WeasyPrint and returns it as ``bytes``.

Convenience helpers
-------------------
:func:`export_docx_to_bytes`
    Like :func:`export_docx` but returns ``bytes`` instead of writing to
    a file path, allowing direct streaming to the browser.

:func:`export_pdf_from_assessment`
    Full pipeline helper: takes an :class:`~ai_gov_gen.assessor.AssessmentResult`,
    generates all artifacts, renders the HTML, and returns PDF ``bytes``.
    Requires an active Flask application context.

:func:`build_export_filename`
    Generates a deterministic, filesystem-safe filename for a given
    system name, export format, and artifact type.

DOCX document structure
-----------------------
The generated ``.docx`` file contains the following sections in order:

1. Cover / metadata table
2. Risk score summary (table)
3. Full AI use policy (one heading + body per section)
4. Governance checklist (table per category)
5. SSP control entries (formatted table per entry)
6. Legal disclaimer

All styles are defined using python-docx's built-in named styles where
possible, falling back to direct paragraph / run formatting when a style
is not available in the default template.

Error handling
--------------
Both exporters raise :class:`ExportError` (a subclass of ``RuntimeError``)
on failure, wrapping the underlying library exception with a descriptive
message that includes the artifact name and failure reason.
"""

from __future__ import annotations

import io
import logging
import re
import unicodedata
from datetime import date
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Custom exception
# ---------------------------------------------------------------------------


class ExportError(RuntimeError):
    """Raised when document export fails.

    Attributes:
        format: The export format that failed (``"docx"`` or ``"pdf"``).
        reason: Human-readable description of the failure cause.
    """

    def __init__(self, format: str, reason: str, cause: BaseException | None = None) -> None:  # noqa: A002
        self.format = format
        self.reason = reason
        self.cause = cause
        super().__init__(f"Export failed [{format.upper()}]: {reason}")


# ---------------------------------------------------------------------------
# Filename helper
# ---------------------------------------------------------------------------


def build_export_filename(
    system_name: str,
    artifact_type: str,
    export_format: str,
    generated_date: str | None = None,
) -> str:
    """Build a deterministic, filesystem-safe export filename.

    The filename follows the pattern::

        ai_gov_<artifact_type>_<sanitised_system_name>_<date>.<ext>

    For example::

        ai_gov_policy_my_ai_system_2024-01-15.docx

    Args:
        system_name: Name of the AI system being assessed.
        artifact_type: Short artifact type label, e.g. ``"policy"``,
            ``"checklist"``, ``"ssp"``, or ``"full"``.
        export_format: File format extension without the leading dot,
            e.g. ``"docx"`` or ``"pdf"``.
        generated_date: Optional ISO date string (``YYYY-MM-DD``).  If
            ``None``, today's date is used.

    Returns:
        A safe filename string including the extension.
    """
    if generated_date is None:
        generated_date = date.today().isoformat()

    # Normalise unicode and replace unsafe characters
    normalised = unicodedata.normalize("NFKD", system_name)
    ascii_name = normalised.encode("ascii", "ignore").decode("ascii")
    safe_name = re.sub(r"[^a-zA-Z0-9]+", "_", ascii_name).strip("_").lower()
    if not safe_name:
        safe_name = "ai_system"
    safe_name = safe_name[:40]  # Truncate to keep filenames reasonable

    ext = export_format.lstrip(".").lower()
    return f"ai_gov_{artifact_type}_{safe_name}_{generated_date}.{ext}"


# ---------------------------------------------------------------------------
# DOCX style helpers
# ---------------------------------------------------------------------------


def _safe_add_heading(
    doc: Any,
    text: str,
    level: int,
) -> Any:
    """Add a heading paragraph, falling back to bold Normal if style missing.

    Args:
        doc: The :class:`docx.Document` instance.
        text: Heading text.
        level: Heading level (1–4).

    Returns:
        The created :class:`docx.text.paragraph.Paragraph`.
    """
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    try:
        para = doc.add_heading(text, level=level)
    except Exception:  # noqa: BLE001
        para = doc.add_paragraph(text)
        run = para.runs[0] if para.runs else para.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)

    # Apply primary colour to top-level headings
    if level <= 2 and para.runs:
        para.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    return para


def _add_cover_table(doc: Any, metadata: dict[str, str]) -> None:
    """Add a formatted cover metadata table to the document.

    Args:
        doc: The :class:`docx.Document` instance.
        metadata: Ordered dict of label → value pairs.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in metadata.items():
        row = table.add_row()
        label_cell = row.cells[0]
        value_cell = row.cells[1]

        # Label cell styling
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        label_cell.width = doc.sections[0].page_width // 4  # ~25% width

        # Value cell styling
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(str(value))
        value_run.font.size = Pt(9)


def _add_risk_summary_table(
    doc: Any,
    assessment_dict: dict[str, Any],
) -> None:
    """Add a risk score summary table to the document.

    Args:
        doc: The :class:`docx.Document` instance.
        assessment_dict: The ``assessment.to_dict()`` plain dictionary.
    """
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    _safe_add_heading(doc, "Risk Score Summary", level=2)

    # Overall row + one row per category
    category_results = assessment_dict.get("category_results", [])
    rows = 1 + len(category_results)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Category", "Risk Level", "Score / 100"]):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "1A3A5C")

    # Overall row
    overall_row = table.add_row().cells
    overall_row[0].paragraphs[0].add_run("OVERALL").bold = True
    overall_row[1].paragraphs[0].add_run(
        assessment_dict.get("overall_risk_level", "")
    ).bold = True
    overall_row[2].paragraphs[0].add_run(
        f"{assessment_dict.get('overall_score', 0.0):.1f}"
    ).bold = True
    _shade_row_by_risk(overall_row, assessment_dict.get("overall_risk_level", "Low"))

    # Per-category rows
    for cr in category_results:
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(cr.get("category_label", ""))
        row_cells[1].paragraphs[0].add_run(cr.get("risk_level", ""))
        row_cells[2].paragraphs[0].add_run(f"{cr.get('normalised_score', 0.0):.1f}")
        for cell in row_cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_paragraph()  # Spacer


def _shade_cell(cell: Any, hex_colour: str) -> None:
    """Apply a solid background fill to a table cell.

    Args:
        cell: python-docx table cell object.
        hex_colour: Six-character hex colour string without the ``#`` prefix.
    """
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]

    tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


_RISK_LEVEL_HEX: dict[str, str] = {
    "Low": "D1E7DD",
    "Medium": "FFF3CD",
    "High": "F8D7DA",
    "Critical": "D3D3D3",
}


def _shade_row_by_risk(cells: Any, risk_level: str) -> None:
    """Shade all cells in a row according to the risk level colour.

    Args:
        cells: Iterable of python-docx cell objects.
        risk_level: Risk level string (Low/Medium/High/Critical).
    """
    hex_colour = _RISK_LEVEL_HEX.get(risk_level, "FFFFFF")
    for cell in cells:
        _shade_cell(cell, hex_colour)


# ---------------------------------------------------------------------------
# DOCX policy sections renderer
# ---------------------------------------------------------------------------


def _render_policy_sections_to_docx(
    doc: Any,
    policy_dict: dict[str, Any],
) -> None:
    """Render all policy document sections into the DOCX document.

    Processes the plain-text section content produced by the generator,
    converting Markdown-like conventions (bold sub-headings, bullet lists,
    blockquotes, tables) into python-docx paragraph formatting.

    Args:
        doc: The :class:`docx.Document` instance.
        policy_dict: The ``policy_document.to_dict()`` plain dictionary.
    """
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    _safe_add_heading(doc, "AI Use Policy", level=1)

    for section in policy_dict.get("sections", []):
        sec_num = section.get("section_number", "")
        sec_title = section.get("title", "")
        sec_content = section.get("content", "")

        # Section heading
        _safe_add_heading(doc, f"{sec_num}  {sec_title}", level=2)

        # Process content line by line
        _process_section_content(doc, sec_content)

        doc.add_paragraph()  # Spacer between sections


def _process_section_content(doc: Any, content: str) -> None:
    """Parse and write Markdown-like section content into the DOCX document.

    Handles:
    * Sub-headings: lines starting and ending with ``**``
    * Bullet lists: lines starting with ``- `` or ``* ``
    * Blockquotes: lines starting with ``> ``
    * Markdown tables: lines starting with ``| ``
    * Regular paragraphs: all other non-empty lines

    Args:
        doc: The :class:`docx.Document` instance.
        content: Plain-text section content string.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    lines = content.split("\n")
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # Filter separator rows
        data_rows = [r for r in table_rows if not all(c.strip().startswith("-") for c in r)]
        if data_rows:
            max_cols = max(len(r) for r in data_rows)
            tbl = doc.add_table(rows=0, cols=max_cols)
            tbl.style = "Table Grid"
            for i, row_cells in enumerate(data_rows):
                tbl_row = tbl.add_row()
                for j in range(max_cols):
                    cell_text = row_cells[j].strip() if j < len(row_cells) else ""
                    p = tbl_row.cells[j].paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.size = Pt(8)
                    if i == 0:
                        run.bold = True
        doc.add_paragraph()  # spacer after table
        in_table = False
        table_rows = []

    for line in lines:
        stripped = line.strip()

        # Markdown table row
        if stripped.startswith("| ") and stripped.endswith(" |"):
            in_table = True
            cells = [c.strip() for c in stripped[1:-1].split(" | ")]
            table_rows.append(cells)
            continue
        else:
            if in_table:
                flush_table()

        # Sub-heading: **Title**
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            inner = stripped[2:-2]
            para = doc.add_paragraph()
            run = para.add_run(inner)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
            continue

        # Bullet list item
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_formatted_run(para, bullet_text)
            continue

        # Blockquote
        if stripped.startswith("> "):
            quote_text = stripped[2:]
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = _cm_to_emu(1.0)
            run = para.add_run(quote_text)
            run.font.size = Pt(9)
            run.italic = True
            run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
            continue

        # Blank line — skip (paragraph break handled by next content)
        if stripped == "":
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        _add_inline_formatted_run(para, stripped)

    # Flush any remaining table
    if in_table:
        flush_table()


def _add_inline_formatted_run(para: Any, text: str) -> None:
    """Add a run to a paragraph, applying inline bold formatting for ``**text**``.

    Splits the text on ``**`` markers and alternates between normal and bold
    runs.  Handles multiple inline bold segments within a single line.

    Args:
        para: The python-docx Paragraph to add runs to.
        text: The text content, potentially containing ``**bold**`` markers.
    """
    from docx.shared import Pt  # type: ignore[import-untyped]

    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.size = Pt(9)
        if i % 2 == 1:  # Odd-index parts are between ** markers → bold
            run.bold = True


def _cm_to_emu(cm: float) -> int:
    """Convert centimetres to English Metric Units (EMU) for python-docx.

    1 cm = 360000 EMU.

    Args:
        cm: Measurement in centimetres.

    Returns:
        Integer value in EMU.
    """
    return int(cm * 360000)


# ---------------------------------------------------------------------------
# DOCX checklist renderer
# ---------------------------------------------------------------------------


def _render_checklist_to_docx(
    doc: Any,
    checklist_dict: dict[str, Any],
) -> None:
    """Render the governance checklist into the DOCX document.

    Adds one table per risk category containing all triggered checklist
    items.  Rows are colour-coded by risk level.

    Args:
        doc: The :class:`docx.Document` instance.
        checklist_dict: The ``checklist.to_dict()`` plain dictionary.
    """
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    _safe_add_heading(doc, "Annex A — Governance Checklist", level=1)

    total_items = checklist_dict.get("total_items", 0)
    total_critical = checklist_dict.get("total_critical_items", 0)
    framework_label = checklist_dict.get("framework_label", "")

    summary_para = doc.add_paragraph()
    summary_para.add_run(
        f"Total items: {total_items}  |  Critical items: {total_critical}  |  "
        f"Framework: {framework_label}  |  "
        f"Generated: {checklist_dict.get('generated_date', '')}  |  "
        f"Overall risk: {checklist_dict.get('overall_risk_level', '')}"
    ).font.size = Pt(9)

    doc.add_paragraph()  # spacer

    for category in checklist_dict.get("categories", []):
        items = category.get("items", [])
        if not items:
            continue

        # Category heading
        cat_label = category.get("category_label", "")
        cat_risk = category.get("risk_level", "Low")
        cat_score = category.get("normalised_score", 0.0)

        _safe_add_heading(doc, f"{cat_label}  —  Risk: {cat_risk}  ({cat_score:.1f}/100)", level=3)

        # Checklist table: Done | Control Ref | Priority | Action | Rationale
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header row
        headers = ["Done", "Control Ref", "Priority", "Required Action", "Rationale"]
        for cell, header_text in zip(table.rows[0].cells, headers):
            run = cell.paragraphs[0].add_run(header_text)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "1A3A5C")

        # Column widths (approximate proportions)
        col_widths_cm = [1.0, 3.0, 2.0, 7.0, 5.0]
        for i, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = _cm_to_emu(width)

        for item in items:
            row_cells = table.add_row().cells
            risk_lvl = item.get("risk_level", "Low")

            # Checkbox placeholder
            row_cells[0].paragraphs[0].add_run("☐").font.size = Pt(10)

            # Control ref
            ref_run = row_cells[1].paragraphs[0].add_run(item.get("control_ref", ""))
            ref_run.font.size = Pt(7)

            # Priority / risk badge
            pri_run = row_cells[2].paragraphs[0].add_run(risk_lvl)
            pri_run.bold = True
            pri_run.font.size = Pt(8)
            _shade_row_by_risk([row_cells[2]], risk_lvl)

            # Action
            action_run = row_cells[3].paragraphs[0].add_run(item.get("action", ""))
            action_run.font.size = Pt(8)
            action_run.bold = True

            # Rationale
            rat_run = row_cells[4].paragraphs[0].add_run(item.get("rationale", ""))
            rat_run.font.size = Pt(7)

        doc.add_paragraph()  # spacer after each category table


# ---------------------------------------------------------------------------
# DOCX SSP entries renderer
# ---------------------------------------------------------------------------


def _render_ssp_entries_to_docx(
    doc: Any,
    ssp_entries_dicts: list[dict[str, Any]],
    system_name: str,
    generated_date: str,
) -> None:
    """Render SSP control entries into the DOCX document.

    Adds an overview summary table followed by a detailed field table
    for each SSP entry.

    Args:
        doc: The :class:`docx.Document` instance.
        ssp_entries_dicts: List of ``entry.to_dict()`` plain dictionaries.
        system_name: Name of the AI system.
        generated_date: ISO date string for the document.
    """
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    _safe_add_heading(doc, "Annex B — SSP Control Entries", level=1)

    intro_para = doc.add_paragraph()
    intro_para.add_run(
        f"The following SSP control entries document the implementation status of key "
        f"AI governance controls for {system_name}. "
        f"Assessment conducted: {generated_date}."
    ).font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # Summary overview table
    if ssp_entries_dicts:
        _safe_add_heading(doc, "Implementation Status Overview", level=3)
        summary_table = doc.add_table(rows=1, cols=5)
        summary_table.style = "Table Grid"

        summary_headers = ["Entry ID", "Control Name", "Family", "Status", "Risk"]
        for cell, h_text in zip(summary_table.rows[0].cells, summary_headers):
            run = cell.paragraphs[0].add_run(h_text)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "1A3A5C")

        for entry in ssp_entries_dicts:
            row = summary_table.add_row().cells
            row[0].paragraphs[0].add_run(entry.get("entry_id", "")).font.size = Pt(7)
            row[1].paragraphs[0].add_run(entry.get("control_name", "")).font.size = Pt(8)
            row[2].paragraphs[0].add_run(entry.get("control_family", "")).font.size = Pt(8)
            status = entry.get("implementation_status", "")
            status_run = row[3].paragraphs[0].add_run(status)
            status_run.font.size = Pt(8)
            status_run.bold = True
            _shade_row_by_risk([row[3]], entry.get("risk_level", "Low"))
            row[4].paragraphs[0].add_run(entry.get("risk_level", "")).font.size = Pt(8)

        doc.add_paragraph()  # spacer

    # Detailed entries
    _safe_add_heading(doc, "Detailed Control Entries", level=3)

    for entry in ssp_entries_dicts:
        entry_id = entry.get("entry_id", "")
        control_name = entry.get("control_name", "")
        impl_status = entry.get("implementation_status", "")
        risk_lvl = entry.get("risk_level", "Low")

        # Entry heading
        entry_para = doc.add_paragraph()
        entry_para.paragraph_format.space_before = _cm_to_emu(0.3)
        id_run = entry_para.add_run(f"{entry_id}  ")
        id_run.bold = True
        id_run.font.size = Pt(10)
        id_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        name_run = entry_para.add_run(control_name)
        name_run.bold = True
        name_run.font.size = Pt(10)

        # Entry detail table
        det_table = doc.add_table(rows=0, cols=2)
        det_table.style = "Table Grid"

        fields = [
            ("Control Family",       entry.get("control_family", "")),
            ("Control Reference",    entry.get("control_ref", "")),
            ("Description",          entry.get("control_description", "")),
            ("Responsible Role",     entry.get("responsible_role", "")),
            ("Implementation Status",impl_status),
            ("Implementation Detail",entry.get("implementation_detail", "")),
            ("Assessment Finding",   entry.get("assessment_finding", "")),
            ("Remediation Action",   entry.get("remediation_action", "")),
            ("Evidence Artefacts",   ", ".join(entry.get("evidence_artifacts", []))),
            ("Applicable Frameworks",", ".join(entry.get("frameworks", []))),
        ]

        for label, value in fields:
            row = det_table.add_row()
            label_cell = row.cells[0]
            value_cell = row.cells[1]

            label_run = label_cell.paragraphs[0].add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

            value_run = value_cell.paragraphs[0].add_run(str(value))
            value_run.font.size = Pt(8)

            if label == "Implementation Status":
                _shade_row_by_risk([label_cell, value_cell], risk_lvl)
                value_run.bold = True
            if label == "Remediation Action" and impl_status not in (
                "Implemented", "Not Applicable"
            ):
                value_run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)

        doc.add_paragraph()  # spacer after each entry


# ---------------------------------------------------------------------------
# DOCX disclaimer and signature block
# ---------------------------------------------------------------------------


def _add_disclaimer(doc: Any) -> None:
    """Add the legal disclaimer paragraph to the DOCX document.

    Args:
        doc: The :class:`docx.Document` instance.
    """
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    doc.add_paragraph()
    _safe_add_heading(doc, "Legal Disclaimer", level=3)

    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run(
        "This document was generated automatically by AI Gov Gen based on questionnaire "
        "responses provided by the system owner. It is intended as a starting point for "
        "governance documentation and does not constitute legal, regulatory, or professional "
        "compliance advice. The organisation is solely responsible for ensuring that its "
        "AI systems comply with applicable laws and regulations. This document must be "
        "reviewed, validated, and approved by qualified compliance, legal, and/or security "
        "professionals before formal adoption or submission to any regulatory body."
    )
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.color.rgb = RGBColor(0x58, 0x15, 0x1C)
    disclaimer_para.paragraph_format.left_indent = _cm_to_emu(0.5)


def _add_signature_block(doc: Any) -> None:
    """Add a document approval signature table to the DOCX document.

    Args:
        doc: The :class:`docx.Document` instance.
    """
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    doc.add_paragraph()
    _safe_add_heading(doc, "Document Approval and Sign-off", level=3)

    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.style = "Table Grid"

    headers = ["Role", "Name & Signature", "Date"]
    for cell, h_text in zip(sig_table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(h_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "1A3A5C")

    roles = [
        "System Owner",
        "AI Governance Authority",
        "Information Security Officer",
        "Legal / Compliance Reviewer",
    ]
    for role in roles:
        row = sig_table.add_row()
        row.cells[0].paragraphs[0].add_run(role).font.size = Pt(9)
        # Leave name/signature and date cells blank for manual completion
        for cell in row.cells[1:]:
            # Add some vertical space for signing
            cell.paragraphs[0].add_run("\n\n")


# ---------------------------------------------------------------------------
# Public DOCX export functions
# ---------------------------------------------------------------------------


def export_docx_to_bytes(
    policy_dict: dict[str, Any],
    assessment_dict: dict[str, Any],
    checklist_dict: dict[str, Any] | None = None,
    ssp_entries_dicts: list[dict[str, Any]] | None = None,
) -> bytes:
    """Export governance artifacts to a DOCX file and return as bytes.

    This function builds a complete python-docx document containing the
    policy document, risk summary, governance checklist, and SSP entries,
    then serialises it to a byte buffer and returns the bytes.

    The function operates entirely on plain dict representations of the
    artifacts so it does not require generator dataclass instances.

    Args:
        policy_dict: The ``policy_document.to_dict()`` plain dictionary.
        assessment_dict: The ``assessment.to_dict()`` plain dictionary.
        checklist_dict: Optional ``checklist.to_dict()`` plain dictionary.
            When ``None``, the checklist annex is omitted.
        ssp_entries_dicts: Optional list of ``entry.to_dict()`` plain dicts.
            When ``None`` or empty, the SSP annex is omitted.

    Returns:
        Raw DOCX bytes ready for streaming or writing to disk.

    Raises:
        ExportError: If document construction or serialisation fails.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Inches, Pt  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ExportError(
            "docx",
            "python-docx is not installed. Run: pip install python-docx",
            exc,
        ) from exc

    system_name = policy_dict.get("system_name", "AI System")
    generated_date = policy_dict.get("generated_date", date.today().isoformat())

    logger.info(
        "Building DOCX export for system '%s'.",
        system_name,
    )

    try:
        doc = Document()

        # ----------------------------------------------------------------
        # Document properties / default font
        # ----------------------------------------------------------------
        core_props = doc.core_properties
        core_props.title = policy_dict.get("title", f"AI Use Policy — {system_name}")
        core_props.author = policy_dict.get("system_owner", "AI Gov Gen")
        core_props.subject = "AI Governance Documentation"
        core_props.keywords = (
            f"AI governance, {policy_dict.get('framework_label', '')}, "
            f"{policy_dict.get('overall_risk_level', '')} risk"
        )
        core_props.category = "Governance"

        # Set default font via Normal style
        try:
            normal_style = doc.styles["Normal"]
            normal_style.font.name = "Calibri"
            normal_style.font.size = Pt(10)
        except Exception:  # noqa: BLE001
            pass

        # ----------------------------------------------------------------
        # Cover / metadata section
        # ----------------------------------------------------------------
        _safe_add_heading(doc, policy_dict.get("title", f"AI Use Policy — {system_name}"), level=1)

        doc.add_paragraph(
            "CLASSIFICATION: INTERNAL — RESTRICTED"
        ).runs[0].bold = True if doc.paragraphs[-1].runs else False
        if doc.paragraphs[-1].runs:
            doc.paragraphs[-1].runs[0].bold = True

        doc.add_paragraph()  # spacer

        metadata = {
            "System Name": system_name,
            "System Owner": policy_dict.get("system_owner", ""),
            "System Purpose": policy_dict.get("system_purpose", ""),
            "Compliance Framework": policy_dict.get("framework_label", ""),
            "Document ID": policy_dict.get("document_id", ""),
            "Version": policy_dict.get("version", "1.0"),
            "Generated Date": generated_date,
            "Overall Risk Level": assessment_dict.get("overall_risk_level", ""),
            "Overall Risk Score": f"{assessment_dict.get('overall_score', 0.0):.1f} / 100",
        }
        _add_cover_table(doc, metadata)

        doc.add_paragraph()  # spacer

        # ----------------------------------------------------------------
        # Warnings (if any)
        # ----------------------------------------------------------------
        warnings = policy_dict.get("warnings", [])
        if warnings:
            _safe_add_heading(doc, "Assessment Warnings", level=3)
            for warning in warnings:
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(warning)
                run.font.size = Pt(9)

        # ----------------------------------------------------------------
        # Risk score summary
        # ----------------------------------------------------------------
        _add_risk_summary_table(doc, assessment_dict)

        # ----------------------------------------------------------------
        # Policy sections
        # ----------------------------------------------------------------
        doc.add_page_break()
        _render_policy_sections_to_docx(doc, policy_dict)

        # ----------------------------------------------------------------
        # Governance checklist (Annex A)
        # ----------------------------------------------------------------
        if checklist_dict:
            doc.add_page_break()
            _render_checklist_to_docx(doc, checklist_dict)

        # ----------------------------------------------------------------
        # SSP entries (Annex B)
        # ----------------------------------------------------------------
        if ssp_entries_dicts:
            doc.add_page_break()
            _render_ssp_entries_to_docx(
                doc, ssp_entries_dicts, system_name, generated_date
            )

        # ----------------------------------------------------------------
        # Disclaimer and signature block
        # ----------------------------------------------------------------
        doc.add_page_break()
        _add_disclaimer(doc)
        doc.add_paragraph()  # spacer
        _add_signature_block(doc)

        # ----------------------------------------------------------------
        # Serialise to bytes
        # ----------------------------------------------------------------
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()

        logger.info(
            "DOCX export complete for '%s' (%d bytes).",
            system_name,
            len(docx_bytes),
        )
        return docx_bytes

    except ExportError:
        raise
    except Exception as exc:
        logger.exception("DOCX export failed for '%s'.", system_name)
        raise ExportError(
            "docx",
            f"Failed to build DOCX for '{system_name}': {exc}",
            exc,
        ) from exc


def export_docx(
    policy_dict: dict[str, Any],
    assessment_dict: dict[str, Any],
    output_path: str | Path,
    checklist_dict: dict[str, Any] | None = None,
    ssp_entries_dicts: list[dict[str, Any]] | None = None,
) -> Path:
    """Export governance artifacts to a DOCX file on disk.

    Args:
        policy_dict: The ``policy_document.to_dict()`` plain dictionary.
        assessment_dict: The ``assessment.to_dict()`` plain dictionary.
        output_path: Filesystem path where the ``.docx`` file will be
            written.  Parent directories are created if they do not exist.
        checklist_dict: Optional ``checklist.to_dict()`` plain dictionary.
        ssp_entries_dicts: Optional list of ``entry.to_dict()`` plain dicts.

    Returns:
        The resolved :class:`pathlib.Path` of the written file.

    Raises:
        ExportError: If document construction or file I/O fails.
    """
    docx_bytes = export_docx_to_bytes(
        policy_dict=policy_dict,
        assessment_dict=assessment_dict,
        checklist_dict=checklist_dict,
        ssp_entries_dicts=ssp_entries_dicts,
    )

    out_path = Path(output_path)
    try:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_bytes(docx_bytes)
        logger.info("DOCX file written to '%s'.", out_path)
        return out_path
    except OSError as exc:
        raise ExportError(
            "docx",
            f"Failed to write DOCX to '{out_path}': {exc}",
            exc,
        ) from exc


# ---------------------------------------------------------------------------
# PDF export helpers
# ---------------------------------------------------------------------------


def _build_pdf_bytes(html_content: str) -> bytes:
    """Convert an HTML string to PDF bytes using WeasyPrint.

    WeasyPrint is imported here (rather than at module level) so that the
    rest of the module can be imported in environments where WeasyPrint
    is not installed or its system-level dependencies are absent.

    Args:
        html_content: Complete HTML document string as produced by
            :func:`~ai_gov_gen.generator.render_policy_html`.

    Returns:
        Raw PDF bytes.

    Raises:
        ExportError: If WeasyPrint cannot be imported or fails during
            rendering.
    """
    try:
        import weasyprint  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ExportError(
            "pdf",
            "WeasyPrint is not installed. Run: pip install weasyprint",
            exc,
        ) from exc

    try:
        logger.debug("Starting WeasyPrint PDF rendering (HTML length: %d chars).", len(html_content))
        pdf_bytes: bytes = weasyprint.HTML(string=html_content).write_pdf()
        logger.debug("WeasyPrint PDF rendering complete (%d bytes).", len(pdf_bytes))
        return pdf_bytes
    except Exception as exc:
        raise ExportError(
            "pdf",
            f"WeasyPrint PDF rendering failed: {exc}",
            exc,
        ) from exc


# ---------------------------------------------------------------------------
# Public PDF export functions
# ---------------------------------------------------------------------------


def export_pdf_from_html(html_content: str) -> bytes:
    """Convert an HTML string to PDF bytes.

    This low-level function accepts a pre-rendered HTML string and returns
    the PDF byte stream.  The HTML is typically produced by
    :func:`~ai_gov_gen.generator.render_policy_html`.

    Args:
        html_content: Complete, self-contained HTML string.

    Returns:
        Raw PDF bytes suitable for streaming to the browser or writing
        to disk.

    Raises:
        ExportError: If WeasyPrint is unavailable or rendering fails.
        ValueError: If ``html_content`` is empty.
    """
    if not html_content or not html_content.strip():
        raise ValueError("html_content must not be empty.")

    return _build_pdf_bytes(html_content)


def export_pdf(
    html_content: str,
    output_path: str | Path,
) -> Path:
    """Convert an HTML string to a PDF file on disk.

    Args:
        html_content: Complete HTML document string.
        output_path: Filesystem path where the ``.pdf`` file will be
            written.  Parent directories are created if they do not exist.

    Returns:
        The resolved :class:`pathlib.Path` of the written file.

    Raises:
        ExportError: If PDF rendering or file I/O fails.
        ValueError: If ``html_content`` is empty.
    """
    pdf_bytes = export_pdf_from_html(html_content)

    out_path = Path(output_path)
    try:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_bytes(pdf_bytes)
        logger.info("PDF file written to '%s'.", out_path)
        return out_path
    except OSError as exc:
        raise ExportError(
            "pdf",
            f"Failed to write PDF to '{out_path}': {exc}",
            exc,
        ) from exc


def export_pdf_from_assessment(
    assessment: Any,
    include_checklist: bool = True,
    include_ssp: bool = True,
) -> bytes:
    """Full pipeline: score → generate → render → PDF.

    Convenience function that generates all governance artifacts from an
    :class:`~ai_gov_gen.assessor.AssessmentResult`, renders the combined
    HTML document via :func:`~ai_gov_gen.generator.render_policy_html`,
    and converts it to PDF bytes using WeasyPrint.

    An active Flask application context is required because
    :func:`~ai_gov_gen.generator.render_policy_html` uses Flask's Jinja2
    template environment.

    Args:
        assessment: A :class:`~ai_gov_gen.assessor.AssessmentResult`
            produced by :func:`~ai_gov_gen.assessor.score_responses`.
        include_checklist: Whether to include the governance checklist
            annex in the PDF (default ``True``).
        include_ssp: Whether to include the SSP control entries annex
            in the PDF (default ``True``).

    Returns:
        Raw PDF bytes.

    Raises:
        ExportError: If artifact generation, HTML rendering, or PDF
            conversion fails.
        RuntimeError: If called outside an active Flask application
            context.
    """
    try:
        from ai_gov_gen.generator import (
            generate_checklist,
            generate_policy_document,
            generate_ssp_entries,
            render_policy_html,
        )
    except ImportError as exc:
        raise ExportError(
            "pdf",
            "ai_gov_gen.generator could not be imported.",
            exc,
        ) from exc

    system_name = getattr(assessment, "system_name", "AI System")

    logger.info(
        "Starting full PDF export pipeline for '%s'.",
        system_name,
    )

    try:
        policy_document = generate_policy_document(assessment)
        checklist = generate_checklist(assessment) if include_checklist else None
        ssp_entries_list = generate_ssp_entries(assessment) if include_ssp else None

        html_content = render_policy_html(
            policy_document=policy_document,
            assessment=assessment,
            checklist=checklist,
            ssp_entries=ssp_entries_list,
        )
    except RuntimeError:
        # Re-raise Flask context errors unchanged
        raise
    except Exception as exc:
        logger.exception(
            "Artifact generation or HTML rendering failed for '%s'.", system_name
        )
        raise ExportError(
            "pdf",
            f"Failed to generate artifacts or render HTML for '{system_name}': {exc}",
            exc,
        ) from exc

    pdf_bytes = _build_pdf_bytes(html_content)

    logger.info(
        "Full PDF export pipeline complete for '%s' (%d bytes).",
        system_name,
        len(pdf_bytes),
    )
    return pdf_bytes


def export_docx_from_assessment(
    assessment: Any,
    include_checklist: bool = True,
    include_ssp: bool = True,
) -> bytes:
    """Full pipeline: score → generate → DOCX bytes.

    Convenience function that generates all governance artifacts from an
    :class:`~ai_gov_gen.assessor.AssessmentResult` and returns DOCX bytes.
    Does not require a Flask application context.

    Args:
        assessment: A :class:`~ai_gov_gen.assessor.AssessmentResult`
            produced by :func:`~ai_gov_gen.assessor.score_responses`.
        include_checklist: Whether to include the governance checklist
            annex in the DOCX (default ``True``).
        include_ssp: Whether to include the SSP control entries annex
            in the DOCX (default ``True``).

    Returns:
        Raw DOCX bytes.

    Raises:
        ExportError: If artifact generation or DOCX construction fails.
    """
    try:
        from ai_gov_gen.generator import (
            generate_checklist,
            generate_policy_document,
            generate_ssp_entries,
        )
    except ImportError as exc:
        raise ExportError(
            "docx",
            "ai_gov_gen.generator could not be imported.",
            exc,
        ) from exc

    system_name = getattr(assessment, "system_name", "AI System")

    logger.info(
        "Starting full DOCX export pipeline for '%s'.",
        system_name,
    )

    try:
        policy_document = generate_policy_document(assessment)
        checklist = generate_checklist(assessment) if include_checklist else None
        ssp_entries_list = generate_ssp_entries(assessment) if include_ssp else None
    except Exception as exc:
        logger.exception(
            "Artifact generation failed for '%s'.", system_name
        )
        raise ExportError(
            "docx",
            f"Failed to generate artifacts for '{system_name}': {exc}",
            exc,
        ) from exc

    return export_docx_to_bytes(
        policy_dict=policy_document.to_dict(),
        assessment_dict=assessment.to_dict(),
        checklist_dict=checklist.to_dict() if checklist else None,
        ssp_entries_dicts=(
            [e.to_dict() for e in ssp_entries_list] if ssp_entries_list else None
        ),
    )
