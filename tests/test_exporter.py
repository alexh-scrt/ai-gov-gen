"""Tests for the DOCX and PDF export engine in ai_gov_gen/exporter.py.

Covers:
* :func:`build_export_filename` — filename generation for various inputs
* :func:`export_docx_to_bytes` — DOCX byte generation from plain dicts
* :func:`export_docx` — DOCX file writing
* :func:`export_pdf_from_html` — PDF byte generation from HTML
* :func:`export_pdf` — PDF file writing
* :func:`export_docx_from_assessment` — full DOCX pipeline
* :func:`export_pdf_from_assessment` — full PDF pipeline (Flask context required)
* :class:`ExportError` — exception structure
* Edge cases: empty inputs, invalid paths, missing dependencies
"""

from __future__ import annotations

import io
import os
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest

from ai_gov_gen.assessor import AssessmentResult, CategoryResult, _map_score_to_risk_level
from ai_gov_gen.exporter import (
    ExportError,
    _build_pdf_bytes,
    build_export_filename,
    export_docx,
    export_docx_from_assessment,
    export_docx_to_bytes,
    export_pdf,
    export_pdf_from_html,
)
from ai_gov_gen.generator import (
    generate_all,
    generate_checklist,
    generate_policy_document,
    generate_ssp_entries,
)
from ai_gov_gen.questions import CATEGORIES


# ---------------------------------------------------------------------------
# Helpers and fixtures
# ---------------------------------------------------------------------------


def _make_category_result(
    cat_id: str,
    score: float = 30.0,
) -> CategoryResult:
    labels = {
        "data": "Data Governance",
        "model": "Model Risk",
        "ops": "Operational Security",
        "compliance": "Compliance & Audit",
    }
    return CategoryResult(
        category_id=cat_id,
        category_label=labels.get(cat_id, cat_id.capitalize()),
        normalised_score=score,
        risk_level=_map_score_to_risk_level(score),
        raw_score=score * 0.3,
        max_possible_score=100.0,
        answered_questions=5,
        total_questions=5,
        question_details=[],
    )


def _make_assessment(
    system_name: str = "Test AI System",
    system_owner: str = "Engineering Team",
    system_purpose: str = "Automates internal document review.",
    framework_id: str = "enterprise",
    overall_score: float = 35.0,
) -> AssessmentResult:
    category_results = [
        _make_category_result("data", 25.0),
        _make_category_result("model", 45.0),
        _make_category_result("ops", 30.0),
        _make_category_result("compliance", 20.0),
    ]
    responses: dict[str, Any] = {
        "meta_01": system_name,
        "meta_02": system_owner,
        "meta_03": system_purpose,
        "meta_04": framework_id,
        "data_02": "yes_manual",
        "data_03": "yes_documented",
        "data_05": "both",
        "data_06": "yes_enforced",
        "model_03": "comprehensive",
        "model_04": "xai_tooling",
        "model_05": "independent",
        "model_08": "mandatory",
        "ops_02": "comprehensive",
        "ops_03": "all_users",
        "ops_05": "ai_specific",
        "comp_02": "dedicated_committee",
        "comp_03": "published_signed",
        "comp_08": "fully_documented",
    }
    framework_labels = {
        "nist_ai_rmf": "NIST AI RMF",
        "hks": "HKS AI Risk Framework",
        "cmmc_l2": "CMMC Level 2",
        "lloyds": "Lloyd's Market AI Guidelines",
        "enterprise": "Enterprise Generic",
    }
    return AssessmentResult(
        category_results=category_results,
        overall_score=overall_score,
        overall_risk_level=_map_score_to_risk_level(overall_score),
        framework_id=framework_id,
        framework_label=framework_labels.get(framework_id, framework_id),
        system_name=system_name,
        system_owner=system_owner,
        system_purpose=system_purpose,
        responses=responses,
        warnings=[],
    )


@pytest.fixture()
def assessment() -> AssessmentResult:
    return _make_assessment()


@pytest.fixture()
def artifacts(assessment: AssessmentResult):
    return generate_all(assessment)


@pytest.fixture()
def policy_dict(assessment: AssessmentResult) -> dict[str, Any]:
    return generate_policy_document(assessment).to_dict()


@pytest.fixture()
def assessment_dict(assessment: AssessmentResult) -> dict[str, Any]:
    return assessment.to_dict()


@pytest.fixture()
def checklist_dict(assessment: AssessmentResult) -> dict[str, Any]:
    return generate_checklist(assessment).to_dict()


@pytest.fixture()
def ssp_entries_dicts(assessment: AssessmentResult) -> list[dict[str, Any]]:
    return [e.to_dict() for e in generate_ssp_entries(assessment)]


@pytest.fixture()
def flask_app(tmp_path: Path):
    from ai_gov_gen import create_app
    return create_app(
        test_config={
            "TESTING": True,
            "OUTPUT_FOLDER": str(tmp_path / "output"),
            "SECRET_KEY": "test-secret",
        }
    )


# ---------------------------------------------------------------------------
# ExportError tests
# ---------------------------------------------------------------------------


class TestExportError:
    """Tests for the custom ExportError exception class."""

    def test_is_runtime_error(self) -> None:
        err = ExportError("docx", "Something went wrong.")
        assert isinstance(err, RuntimeError)

    def test_format_attribute(self) -> None:
        err = ExportError("pdf", "Render failed.")
        assert err.format == "pdf"

    def test_reason_attribute(self) -> None:
        err = ExportError("docx", "Build failed.")
        assert err.reason == "Build failed."

    def test_cause_attribute_none_by_default(self) -> None:
        err = ExportError("docx", "Error.")
        assert err.cause is None

    def test_cause_attribute_stored(self) -> None:
        original = ValueError("Original error.")
        err = ExportError("pdf", "Wrapped error.", cause=original)
        assert err.cause is original

    def test_str_includes_format_and_reason(self) -> None:
        err = ExportError("docx", "Template failed.")
        msg = str(err)
        assert "DOCX" in msg
        assert "Template failed." in msg

    def test_can_be_raised_and_caught(self) -> None:
        with pytest.raises(ExportError, match="Export failed"):
            raise ExportError("pdf", "Test error.")

    def test_caught_as_runtime_error(self) -> None:
        with pytest.raises(RuntimeError):
            raise ExportError("docx", "Test.")


# ---------------------------------------------------------------------------
# build_export_filename tests
# ---------------------------------------------------------------------------


class TestBuildExportFilename:
    """Tests for build_export_filename()."""

    def test_returns_string(self) -> None:
        result = build_export_filename("My AI", "policy", "docx")
        assert isinstance(result, str)

    def test_has_correct_extension_docx(self) -> None:
        result = build_export_filename("My AI", "policy", "docx")
        assert result.endswith(".docx")

    def test_has_correct_extension_pdf(self) -> None:
        result = build_export_filename("My AI", "policy", "pdf")
        assert result.endswith(".pdf")

    def test_includes_artifact_type(self) -> None:
        result = build_export_filename("My AI", "checklist", "docx")
        assert "checklist" in result

    def test_includes_sanitised_system_name(self) -> None:
        result = build_export_filename("My AI System", "policy", "docx")
        assert "my" in result.lower() or "ai" in result.lower()

    def test_includes_generated_date_when_provided(self) -> None:
        result = build_export_filename("Test AI", "policy", "docx", generated_date="2024-06-15")
        assert "2024-06-15" in result

    def test_uses_today_when_date_not_provided(self) -> None:
        from datetime import date
        today = date.today().isoformat()
        result = build_export_filename("Test AI", "policy", "docx")
        assert today in result

    def test_special_characters_removed(self) -> None:
        result = build_export_filename("AI/System!@#$", "policy", "pdf")
        assert "/" not in result
        assert "!" not in result
        assert "@" not in result
        assert "#" not in result

    def test_spaces_replaced_with_underscores(self) -> None:
        result = build_export_filename("My AI System", "policy", "docx")
        # Spaces in system name should be replaced
        assert " " not in result

    def test_empty_system_name_uses_fallback(self) -> None:
        result = build_export_filename("", "policy", "docx")
        assert "ai_system" in result or "policy" in result

    def test_very_long_system_name_truncated(self) -> None:
        long_name = "A" * 200
        result = build_export_filename(long_name, "policy", "docx")
        # Should not produce an excessively long filename
        assert len(result) < 150

    def test_starts_with_ai_gov_prefix(self) -> None:
        result = build_export_filename("Test AI", "policy", "docx")
        assert result.startswith("ai_gov_")

    def test_unicode_system_name_handled(self) -> None:
        result = build_export_filename("AI Système Générique", "policy", "docx")
        assert isinstance(result, str)
        assert result.endswith(".docx")

    def test_dot_in_extension_stripped(self) -> None:
        result = build_export_filename("Test AI", "policy", ".docx")
        assert result.endswith(".docx")
        assert ".." not in result

    @pytest.mark.parametrize(
        "artifact_type",
        ["policy", "checklist", "ssp", "full"],
    )
    def test_all_artifact_types_produce_valid_filename(
        self, artifact_type: str
    ) -> None:
        result = build_export_filename("Test AI", artifact_type, "docx")
        assert isinstance(result, str) and len(result) > 0
        assert artifact_type in result


# ---------------------------------------------------------------------------
# export_docx_to_bytes tests
# ---------------------------------------------------------------------------


class TestExportDocxToBytes:
    """Tests for export_docx_to_bytes()."""

    def test_returns_bytes(self, policy_dict, assessment_dict) -> None:
        result = export_docx_to_bytes(policy_dict, assessment_dict)
        assert isinstance(result, bytes)

    def test_returns_non_empty_bytes(self, policy_dict, assessment_dict) -> None:
        result = export_docx_to_bytes(policy_dict, assessment_dict)
        assert len(result) > 0

    def test_produces_valid_docx_magic_bytes(self, policy_dict, assessment_dict) -> None:
        """DOCX files are ZIP archives and start with the ZIP magic bytes PK."""
        result = export_docx_to_bytes(policy_dict, assessment_dict)
        # ZIP/DOCX magic bytes: 50 4B 03 04
        assert result[:2] == b"PK"

    def test_docx_can_be_reopened_by_python_docx(
        self, policy_dict, assessment_dict
    ) -> None:
        """The generated bytes should be a valid DOCX that python-docx can open."""
        from docx import Document

        result = export_docx_to_bytes(policy_dict, assessment_dict)
        doc = Document(io.BytesIO(result))
        # Should have at least one paragraph
        assert len(doc.paragraphs) > 0

    def test_docx_contains_system_name(
        self, policy_dict, assessment_dict
    ) -> None:
        """The system name should appear in the document text."""
        from docx import Document

        result = export_docx_to_bytes(policy_dict, assessment_dict)
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        system_name = policy_dict.get("system_name", "")
        assert system_name in all_text

    def test_docx_contains_system_owner(
        self, policy_dict, assessment_dict
    ) -> None:
        from docx import Document

        result = export_docx_to_bytes(policy_dict, assessment_dict)
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        system_owner = policy_dict.get("system_owner", "")
        assert system_owner in all_text

    def test_docx_with_checklist_is_larger(
        self, policy_dict, assessment_dict, checklist_dict
    ) -> None:
        without_checklist = export_docx_to_bytes(policy_dict, assessment_dict)
        with_checklist = export_docx_to_bytes(
            policy_dict, assessment_dict, checklist_dict=checklist_dict
        )
        assert len(with_checklist) >= len(without_checklist)

    def test_docx_with_ssp_is_larger(
        self, policy_dict, assessment_dict, ssp_entries_dicts
    ) -> None:
        without_ssp = export_docx_to_bytes(policy_dict, assessment_dict)
        with_ssp = export_docx_to_bytes(
            policy_dict, assessment_dict, ssp_entries_dicts=ssp_entries_dicts
        )
        assert len(with_ssp) >= len(without_ssp)

    def test_docx_with_all_artifacts(
        self, policy_dict, assessment_dict, checklist_dict, ssp_entries_dicts
    ) -> None:
        result = export_docx_to_bytes(
            policy_dict,
            assessment_dict,
            checklist_dict=checklist_dict,
            ssp_entries_dicts=ssp_entries_dicts,
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_docx_none_checklist_and_ssp_omitted(
        self, policy_dict, assessment_dict
    ) -> None:
        """Passing None for checklist and SSP should not raise."""
        result = export_docx_to_bytes(
            policy_dict,
            assessment_dict,
            checklist_dict=None,
            ssp_entries_dicts=None,
        )
        assert isinstance(result, bytes)

    def test_docx_contains_risk_level(
        self, policy_dict, assessment_dict
    ) -> None:
        from docx import Document

        result = export_docx_to_bytes(policy_dict, assessment_dict)
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        risk_level = assessment_dict.get("overall_risk_level", "")
        assert risk_level in all_text

    def test_docx_has_tables(
        self, policy_dict, assessment_dict
    ) -> None:
        """Generated DOCX should have at least the metadata and risk tables."""
        from docx import Document

        result = export_docx_to_bytes(policy_dict, assessment_dict)
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) >= 1

    def test_idempotent_same_bytes_on_same_input(
        self, policy_dict, assessment_dict
    ) -> None:
        """Two calls with the same input should produce identical byte sizes."""
        result1 = export_docx_to_bytes(policy_dict, assessment_dict)
        result2 = export_docx_to_bytes(policy_dict, assessment_dict)
        # DOCX files include timestamps, so exact bytes may differ,
        # but file sizes should be in the same ballpark (within 10%).
        ratio = len(result1) / max(len(result2), 1)
        assert 0.9 <= ratio <= 1.1

    def test_docx_contains_framework_label(
        self, policy_dict, assessment_dict
    ) -> None:
        from docx import Document

        result = export_docx_to_bytes(policy_dict, assessment_dict)
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        fw_label = policy_dict.get("framework_label", "")
        assert fw_label in all_text

    def test_checklist_items_appear_in_docx(
        self, policy_dict, assessment_dict, checklist_dict
    ) -> None:
        from docx import Document

        result = export_docx_to_bytes(
            policy_dict, assessment_dict, checklist_dict=checklist_dict
        )
        doc = Document(io.BytesIO(result))
        # The checklist table should have been added
        assert len(doc.tables) >= 2  # at least cover table + checklist table

    def test_ssp_entries_appear_in_docx(
        self, policy_dict, assessment_dict, ssp_entries_dicts
    ) -> None:
        from docx import Document

        result = export_docx_to_bytes(
            policy_dict, assessment_dict, ssp_entries_dicts=ssp_entries_dicts
        )
        doc = Document(io.BytesIO(result))
        # Check at least one SSP entry ID appears in the text
        all_text = " ".join(p.text for p in doc.paragraphs)
        if ssp_entries_dicts:
            first_id = ssp_entries_dicts[0].get("entry_id", "")
            assert first_id in all_text

    def test_warnings_appear_in_docx_when_present(
        self, assessment_dict
    ) -> None:
        from docx import Document

        assessment = _make_assessment()
        policy = generate_policy_document(assessment)
        policy_d = policy.to_dict()
        policy_d["warnings"] = ["Critical: Review data provenance immediately."]

        result = export_docx_to_bytes(policy_d, assessment_dict)
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Critical: Review data provenance immediately." in all_text


# ---------------------------------------------------------------------------
# export_docx tests
# ---------------------------------------------------------------------------


class TestExportDocx:
    """Tests for export_docx() file writing."""

    def test_writes_file(self, tmp_path, policy_dict, assessment_dict) -> None:
        output = tmp_path / "test_policy.docx"
        export_docx(policy_dict, assessment_dict, output)
        assert output.exists()

    def test_returns_path_object(self, tmp_path, policy_dict, assessment_dict) -> None:
        output = tmp_path / "test.docx"
        result = export_docx(policy_dict, assessment_dict, output)
        assert isinstance(result, Path)

    def test_returned_path_matches_output(self, tmp_path, policy_dict, assessment_dict) -> None:
        output = tmp_path / "test.docx"
        result = export_docx(policy_dict, assessment_dict, output)
        assert result.resolve() == output.resolve()

    def test_file_is_non_empty(self, tmp_path, policy_dict, assessment_dict) -> None:
        output = tmp_path / "test.docx"
        export_docx(policy_dict, assessment_dict, output)
        assert output.stat().st_size > 0

    def test_creates_parent_dirs(self, tmp_path, policy_dict, assessment_dict) -> None:
        output = tmp_path / "subdir" / "nested" / "test.docx"
        export_docx(policy_dict, assessment_dict, output)
        assert output.exists()

    def test_accepts_string_path(self, tmp_path, policy_dict, assessment_dict) -> None:
        output_str = str(tmp_path / "test.docx")
        result = export_docx(policy_dict, assessment_dict, output_str)
        assert Path(output_str).exists()

    def test_with_checklist(self, tmp_path, policy_dict, assessment_dict, checklist_dict) -> None:
        output = tmp_path / "with_checklist.docx"
        export_docx(policy_dict, assessment_dict, output, checklist_dict=checklist_dict)
        assert output.exists()
        assert output.stat().st_size > 0

    def test_with_all_artifacts(
        self, tmp_path, policy_dict, assessment_dict, checklist_dict, ssp_entries_dicts
    ) -> None:
        output = tmp_path / "full.docx"
        export_docx(
            policy_dict,
            assessment_dict,
            output,
            checklist_dict=checklist_dict,
            ssp_entries_dicts=ssp_entries_dicts,
        )
        assert output.exists()
        assert output.stat().st_size > 0


# ---------------------------------------------------------------------------
# export_pdf_from_html tests
# ---------------------------------------------------------------------------


class TestExportPdfFromHtml:
    """Tests for export_pdf_from_html() using WeasyPrint."""

    _MINIMAL_HTML = (
        "<!DOCTYPE html><html><head><title>Test</title></head>"
        "<body><h1>Test PDF</h1><p>Content.</p></body></html>"
    )

    def test_returns_bytes(self) -> None:
        result = export_pdf_from_html(self._MINIMAL_HTML)
        assert isinstance(result, bytes)

    def test_returns_non_empty_bytes(self) -> None:
        result = export_pdf_from_html(self._MINIMAL_HTML)
        assert len(result) > 0

    def test_pdf_magic_bytes(self) -> None:
        """PDF files start with the byte sequence %%PDF."""
        result = export_pdf_from_html(self._MINIMAL_HTML)
        assert result[:4] == b"%PDF"

    def test_empty_html_raises_value_error(self) -> None:
        with pytest.raises(ValueError):
            export_pdf_from_html("")

    def test_whitespace_only_html_raises_value_error(self) -> None:
        with pytest.raises(ValueError):
            export_pdf_from_html("   \n\t  ")

    def test_larger_html_produces_larger_pdf(self) -> None:
        small = export_pdf_from_html(self._MINIMAL_HTML)
        large_html = (
            "<!DOCTYPE html><html><head><title>Big</title></head><body>"
            + "<p>Content paragraph.</p>" * 200
            + "</body></html>"
        )
        large = export_pdf_from_html(large_html)
        assert len(large) >= len(small)

    def test_complex_html_produces_valid_pdf(
        self, flask_app, assessment: AssessmentResult
    ) -> None:
        """A full policy HTML document should convert to a valid PDF."""
        from ai_gov_gen.generator import render_policy_html

        artifacts = generate_all(assessment)
        with flask_app.app_context():
            html_content = render_policy_html(
                artifacts.policy_document,
                artifacts.assessment,
                checklist=artifacts.checklist,
                ssp_entries=artifacts.ssp_entries,
            )

        pdf_bytes = export_pdf_from_html(html_content)
        assert isinstance(pdf_bytes, bytes)
        assert pdf_bytes[:4] == b"%PDF"
        assert len(pdf_bytes) > 10_000  # Policy PDFs should be at least 10 KB


# ---------------------------------------------------------------------------
# export_pdf tests
# ---------------------------------------------------------------------------


class TestExportPdf:
    """Tests for export_pdf() file writing."""

    _MINIMAL_HTML = (
        "<!DOCTYPE html><html><head><title>Test</title></head>"
        "<body><h1>Test PDF</h1><p>Content.</p></body></html>"
    )

    def test_writes_file(self, tmp_path: Path) -> None:
        output = tmp_path / "test.pdf"
        export_pdf(self._MINIMAL_HTML, output)
        assert output.exists()

    def test_returns_path_object(self, tmp_path: Path) -> None:
        output = tmp_path / "test.pdf"
        result = export_pdf(self._MINIMAL_HTML, output)
        assert isinstance(result, Path)

    def test_returned_path_matches_output(self, tmp_path: Path) -> None:
        output = tmp_path / "test.pdf"
        result = export_pdf(self._MINIMAL_HTML, output)
        assert result.resolve() == output.resolve()

    def test_file_is_non_empty(self, tmp_path: Path) -> None:
        output = tmp_path / "test.pdf"
        export_pdf(self._MINIMAL_HTML, output)
        assert output.stat().st_size > 0

    def test_file_starts_with_pdf_magic(self, tmp_path: Path) -> None:
        output = tmp_path / "test.pdf"
        export_pdf(self._MINIMAL_HTML, output)
        assert output.read_bytes()[:4] == b"%PDF"

    def test_creates_parent_dirs(self, tmp_path: Path) -> None:
        output = tmp_path / "subdir" / "nested" / "test.pdf"
        export_pdf(self._MINIMAL_HTML, output)
        assert output.exists()

    def test_accepts_string_path(self, tmp_path: Path) -> None:
        output_str = str(tmp_path / "test.pdf")
        export_pdf(self._MINIMAL_HTML, output_str)
        assert Path(output_str).exists()


# ---------------------------------------------------------------------------
# export_docx_from_assessment tests
# ---------------------------------------------------------------------------


class TestExportDocxFromAssessment:
    """Tests for the full pipeline export_docx_from_assessment()."""

    def test_returns_bytes(self, assessment: AssessmentResult) -> None:
        result = export_docx_from_assessment(assessment)
        assert isinstance(result, bytes)

    def test_returns_non_empty_bytes(self, assessment: AssessmentResult) -> None:
        result = export_docx_from_assessment(assessment)
        assert len(result) > 0

    def test_produces_valid_docx(self, assessment: AssessmentResult) -> None:
        from docx import Document

        result = export_docx_from_assessment(assessment)
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_system_name_in_docx(self, assessment: AssessmentResult) -> None:
        from docx import Document

        result = export_docx_from_assessment(assessment)
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert assessment.system_name in all_text

    def test_without_checklist_is_smaller(
        self, assessment: AssessmentResult
    ) -> None:
        with_checklist = export_docx_from_assessment(assessment, include_checklist=True)
        without_checklist = export_docx_from_assessment(assessment, include_checklist=False)
        assert len(with_checklist) >= len(without_checklist)

    def test_without_ssp_is_smaller(self, assessment: AssessmentResult) -> None:
        with_ssp = export_docx_from_assessment(assessment, include_ssp=True)
        without_ssp = export_docx_from_assessment(assessment, include_ssp=False)
        assert len(with_ssp) >= len(without_ssp)

    def test_all_frameworks_produce_valid_docx(
        self,
    ) -> None:
        for fw in ["nist_ai_rmf", "hks", "cmmc_l2", "lloyds", "enterprise"]:
            assessment = _make_assessment(framework_id=fw)
            result = export_docx_from_assessment(assessment)
            assert isinstance(result, bytes)
            assert result[:2] == b"PK"

    def test_large_system_name_handled(self) -> None:
        assessment = _make_assessment(
            system_name="Very Long AI System Name That Exceeds Normal Limits " * 3
        )
        result = export_docx_from_assessment(assessment)
        assert isinstance(result, bytes) and len(result) > 0

    def test_special_characters_in_system_name_handled(self) -> None:
        assessment = _make_assessment(system_name="AI & Analytics <System> v2.0")
        result = export_docx_from_assessment(assessment)
        assert isinstance(result, bytes) and len(result) > 0

    def test_assessment_with_warnings_handled(
        self,
    ) -> None:
        from docx import Document

        assessment = _make_assessment()
        assessment.warnings.append("Test warning from assessment.")
        result = export_docx_from_assessment(assessment)
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Test warning from assessment." in all_text


# ---------------------------------------------------------------------------
# Full PDF pipeline tests (require Flask context)
# ---------------------------------------------------------------------------


class TestExportPdfFromAssessment:
    """Tests for export_pdf_from_assessment() which requires a Flask context."""

    def test_returns_bytes_in_flask_context(
        self, flask_app, assessment: AssessmentResult
    ) -> None:
        with flask_app.app_context():
            result = export_pdf_from_assessment(assessment)
        assert isinstance(result, bytes)

    def test_returns_non_empty_bytes(
        self, flask_app, assessment: AssessmentResult
    ) -> None:
        with flask_app.app_context():
            result = export_pdf_from_assessment(assessment)
        assert len(result) > 0

    def test_produces_pdf_magic_bytes(
        self, flask_app, assessment: AssessmentResult
    ) -> None:
        with flask_app.app_context():
            result = export_pdf_from_assessment(assessment)
        assert result[:4] == b"%PDF"

    def test_pdf_is_reasonably_sized(
        self, flask_app, assessment: AssessmentResult
    ) -> None:
        with flask_app.app_context():
            result = export_pdf_from_assessment(assessment)
        # A full policy PDF with checklist and SSP should be at least 20 KB
        assert len(result) > 20_000

    def test_without_checklist(
        self, flask_app, assessment: AssessmentResult
    ) -> None:
        with flask_app.app_context():
            result = export_pdf_from_assessment(assessment, include_checklist=False)
        assert result[:4] == b"%PDF"

    def test_without_ssp(
        self, flask_app, assessment: AssessmentResult
    ) -> None:
        with flask_app.app_context():
            result = export_pdf_from_assessment(assessment, include_ssp=False)
        assert result[:4] == b"%PDF"

    def test_without_checklist_and_ssp(
        self, flask_app, assessment: AssessmentResult
    ) -> None:
        with flask_app.app_context():
            result = export_pdf_from_assessment(
                assessment, include_checklist=False, include_ssp=False
            )
        assert result[:4] == b"%PDF"

    def test_raises_runtime_error_without_flask_context(
        self, assessment: AssessmentResult
    ) -> None:
        with pytest.raises(RuntimeError):
            export_pdf_from_assessment(assessment)

    @pytest.mark.parametrize(
        "framework_id",
        ["nist_ai_rmf", "hks", "cmmc_l2", "lloyds", "enterprise"],
    )
    def test_all_frameworks_produce_valid_pdf(
        self, flask_app, framework_id: str
    ) -> None:
        assessment = _make_assessment(framework_id=framework_id)
        with flask_app.app_context():
            result = export_pdf_from_assessment(assessment)
        assert isinstance(result, bytes)
        assert result[:4] == b"%PDF"


# ---------------------------------------------------------------------------
# Integration: filename + export together
# ---------------------------------------------------------------------------


class TestFilenameAndExportIntegration:
    """Integration tests combining filename generation with export."""

    def test_docx_file_with_generated_filename(
        self, tmp_path: Path, policy_dict, assessment_dict
    ) -> None:
        filename = build_export_filename(
            policy_dict.get("system_name", "test"), "policy", "docx"
        )
        output = tmp_path / filename
        export_docx(policy_dict, assessment_dict, output)
        assert output.exists()
        assert output.suffix == ".docx"

    def test_pdf_file_with_generated_filename(
        self, tmp_path: Path
    ) -> None:
        filename = build_export_filename("My Test AI", "policy", "pdf")
        output = tmp_path / filename
        minimal_html = (
            "<!DOCTYPE html><html><head><title>T</title></head>"
            "<body><p>Test.</p></body></html>"
        )
        export_pdf(minimal_html, output)
        assert output.exists()
        assert output.suffix == ".pdf"

    def test_multiple_formats_same_assessment(
        self, tmp_path: Path, policy_dict, assessment_dict
    ) -> None:
        """Same assessment should produce both DOCX and PDF without errors."""
        docx_filename = build_export_filename(
            policy_dict.get("system_name", "test"), "policy", "docx"
        )
        docx_path = export_docx(
            policy_dict, assessment_dict, tmp_path / docx_filename
        )
        assert docx_path.exists()
        assert docx_path.stat().st_size > 0

    def test_export_docx_from_assessment_and_write_to_disk(
        self, tmp_path: Path, assessment: AssessmentResult
    ) -> None:
        docx_bytes = export_docx_from_assessment(assessment)
        filename = build_export_filename(
            assessment.system_name, "full", "docx"
        )
        output = tmp_path / filename
        output.write_bytes(docx_bytes)
        assert output.exists()
        assert output.stat().st_size == len(docx_bytes)


# ---------------------------------------------------------------------------
# Edge case and robustness tests
# ---------------------------------------------------------------------------


class TestEdgeCases:
    """Edge case and robustness tests for the exporter."""

    def test_empty_checklist_dict_handled(
        self, policy_dict, assessment_dict
    ) -> None:
        """An empty categories list in the checklist dict should not raise."""
        empty_checklist = {
            "system_name": "Test AI",
            "system_owner": "Team",
            "framework_id": "enterprise",
            "framework_label": "Enterprise Generic",
            "overall_risk_level": "Low",
            "overall_score": 10.0,
            "generated_date": "2024-01-01",
            "total_items": 0,
            "total_critical_items": 0,
            "warnings": [],
            "categories": [],
        }
        result = export_docx_to_bytes(
            policy_dict, assessment_dict, checklist_dict=empty_checklist
        )
        assert isinstance(result, bytes) and len(result) > 0

    def test_empty_ssp_list_handled(
        self, policy_dict, assessment_dict
    ) -> None:
        result = export_docx_to_bytes(
            policy_dict, assessment_dict, ssp_entries_dicts=[]
        )
        assert isinstance(result, bytes) and len(result) > 0

    def test_policy_with_many_sections(
        self, assessment_dict
    ) -> None:
        """Policy with many sections should still produce a valid DOCX."""
        assessment = _make_assessment()
        policy_doc = generate_policy_document(assessment)
        policy_d = policy_doc.to_dict()

        result = export_docx_to_bytes(policy_d, assessment_dict)
        assert isinstance(result, bytes) and len(result) > 0

    def test_assessment_dict_with_missing_category_results(
        self, policy_dict
    ) -> None:
        """assessment_dict without category_results should not crash."""
        minimal_assessment = {
            "overall_score": 50.0,
            "overall_risk_level": "High",
            "overall_risk_level_index": 2,
            "framework_id": "enterprise",
            "framework_label": "Enterprise Generic",
            "system_name": "Test AI",
            "system_owner": "Test Team",
            "system_purpose": "Test purpose.",
            "warnings": [],
            "category_results": [],
            "highest_risk_category": None,
            "critical_categories": [],
            "high_or_critical_categories": [],
        }
        result = export_docx_to_bytes(policy_dict, minimal_assessment)
        assert isinstance(result, bytes) and len(result) > 0

    def test_policy_section_with_multiline_content(
        self, assessment_dict
    ) -> None:
        """Multi-line section content should be handled without errors."""
        policy_d = {
            "title": "AI Use Policy — Test",
            "document_id": "AI-POL-TEST-001",
            "version": "1.0",
            "system_name": "Test AI",
            "system_owner": "Test Team",
            "system_purpose": "Test purpose.",
            "framework_id": "enterprise",
            "framework_label": "Enterprise Generic",
            "overall_risk_level": "Medium",
            "generated_date": "2024-01-01",
            "warnings": [],
            "sections": [
                {
                    "section_number": "1",
                    "title": "Test Section",
                    "content": (
                        "First paragraph.\n"
                        "\n"
                        "Second paragraph.\n"
                        "- Bullet item one\n"
                        "- Bullet item two\n"
                        "\n"
                        "**Sub-heading**\n"
                        "\n"
                        "> Blockquote content here.\n"
                        "\n"
                        "| Column A | Column B |\n"
                        "|---|---|\n"
                        "| Value 1 | Value 2 |\n"
                    ),
                }
            ],
        }
        result = export_docx_to_bytes(policy_d, assessment_dict)
        from docx import Document
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "First paragraph." in all_text
        assert "Bullet item one" in all_text
        assert "Sub-heading" in all_text

    def test_unicode_content_in_policy(
        self, assessment_dict
    ) -> None:
        """Unicode characters in policy content should be handled."""
        policy_d = {
            "title": "AI Use Policy — Système IA",
            "document_id": "AI-POL-SYS-001",
            "version": "1.0",
            "system_name": "Système IA",
            "system_owner": "Équipe Technique",
            "system_purpose": "Analyse de données multilingues.",
            "framework_id": "enterprise",
            "framework_label": "Enterprise Generic",
            "overall_risk_level": "Low",
            "generated_date": "2024-01-01",
            "warnings": [],
            "sections": [
                {
                    "section_number": "1",
                    "title": "Objet et Portée",
                    "content": "Ce système traite des données multilingues: çàéèü.",
                }
            ],
        }
        result = export_docx_to_bytes(policy_d, assessment_dict)
        assert isinstance(result, bytes) and len(result) > 0
